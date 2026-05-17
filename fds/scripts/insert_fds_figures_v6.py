"""
Restructure manuscript v4 -> v6: move all FDS figures into the main body.

Insertion points in main body (v4 indices):
  - Fig 12 (centreline section) after para 159 ("Critically: FM Global...")
  - Section 4.7 title update: para 178 -> "FDS CFD Simulation"
  - Figs 14 replacement: para 187 (image), 188 (title), 189 (caption body)
  - Figs 15-20 (ceiling HF/H2, H2 BZ, fire section, temp BZ, sensor)
    inserted after para 189 in Section 4.7

Supplementary section (para 309-311): image removed, caption updated to cross-ref.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MANUSCRIPT = (r'G:\My Drive\SAI\Projects\EQIX_SG4-4A_NFPA855_HMA_Waiver_Report'
              r'\__Mickey7_BESS_Fire_Safety_Paper\BESS_Fire_Safety_Paper_Q1_revised_v4.docx')
OUTPUT = (r'G:\My Drive\SAI\Projects\EQIX_SG4-4A_NFPA855_HMA_Waiver_Report'
          r'\__Mickey7_BESS_Fire_Safety_Paper\BESS_Fire_Safety_Paper_Q1_revised_v6.docx')
FIGS_DIR = r'C:\FDS_runs\eqix_preview\figures'


# =============================================================================
# LOW-LEVEL HELPERS
# =============================================================================

def clear_runs(para):
    """Remove all child elements except w:pPr."""
    to_remove = [c for c in para._p if c.tag != qn('w:pPr')]
    for c in to_remove:
        para._p.remove(c)


def set_bold(para, text, size_pt=9):
    clear_runs(para)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)


def set_italic(para, text, size_pt=9):
    clear_runs(para)
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(size_pt)


def set_plain(para, text, size_pt=None):
    clear_runs(para)
    run = para.add_run(text)
    if size_pt:
        run.font.size = Pt(size_pt)


def replace_with_image(para, img_path, width_cm=15.0):
    """Replace paragraph content with a centered image."""
    clear_runs(para)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(img_path, width=Cm(width_cm))


def new_para_after(ref_para):
    """Insert a blank paragraph immediately after ref_para; return it."""
    new_p = OxmlElement('w:p')
    ref_para._p.addnext(new_p)
    return ref_para.__class__(new_p, ref_para._parent)


def insert_image_after(ref_para, img_path, width_cm=15.0):
    p = new_para_after(ref_para)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Cm(width_cm))
    return p


def insert_bold_after(ref_para, text, size_pt=9, center=False):
    p = new_para_after(ref_para)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def insert_italic_after(ref_para, text, size_pt=9):
    p = new_para_after(ref_para)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(size_pt)
    return p


def insert_blank_after(ref_para):
    return new_para_after(ref_para)


def insert_figure_block_after(ref_para, img_path, title, caption, width_cm=15.0):
    """Insert spacer + image + bold title + italic caption after ref_para.
    Returns the caption paragraph (last inserted) for chaining."""
    p0 = insert_blank_after(ref_para)          # spacer before
    p1 = insert_image_after(p0, img_path, width_cm)
    p2 = insert_bold_after(p1, title, center=True)
    p3 = insert_italic_after(p2, caption)
    return p3


# =============================================================================
# FIGURE CONTENT
# =============================================================================
FIGS = {
    'fig12': {
        'file':    'Fig12_HF_Section_PBY150.png',
        'title':   'Figure 12: FDS CFD Simulation -- HF Concentration and Temperature, '
                   'South Aisle Vertical Section (Y = 1.50 m)',
        'caption': ('Figure 12. FDS simulation (NIST FDS 6.10.1, 0.15 m mesh) -- HF mass fraction '
                    '(colour fill, ppm) and temperature (red-orange contours, deg C) at Y = 1.50 m '
                    '(south aisle cross-section, 0.5 m in front of CAB-4 fire face at Y = 2.00 m). '
                    'Four time snapshots: t = 150, 300, 450, 600 s. '
                    'Temperature contours at 50, 100, 200, 400 deg C. '
                    'Dashed blue = 3 ppm ACGIH STEL; dashed red = 30 ppm IDLH. '
                    'Dotted line at Z = 1.50 m (breathing zone). Cabinet cross-sections at X positions shown. '
                    'EQIX SG4-4A BESS Room 1, 9 ACH Stage 0 ventilation. '
                    'Cabinet row: Y = 2.00-2.90 m (centred, 2 m south clearance).'),
    },
    'fig14': {
        'file':    'Fig14_HF_BZ_Plan_PBZ150.png',
        'title':   'Figure 14: FDS CFD Simulation -- HF Concentration at Breathing Zone (Z = 1.50 m)',
        'caption': ('Figure 14. FDS simulation (NIST FDS 6.10.1, 0.15 m mesh, 600 s) -- HF mass '
                    'fraction at breathing zone height Z = 1.50 m, five time snapshots '
                    '(t = 120, 240, 360, 480, 600 s). Colour fill: HF concentration (ppm); '
                    'dashed blue = 3 ppm ACGIH STEL; dashed red = 30 ppm IDLH. '
                    'Cabinet outlines C1-C7 shown; CAB-4 (TR source) in red. '
                    'Ventilation: 9 ACH Stage 0 (supply north wall OPEN, exhaust grilles at ceiling). '
                    'EQIX SG4-4A BESS Room 1, NMC Galaxy LBF 242.76 kWh.'),
    },
    'fig15': {
        'file':    'Temp_BZ_Plan_PBZ150.png',
        'title':   'Figure 15: FDS CFD Simulation -- Temperature at Breathing Zone (Z = 1.50 m, Plan View)',
        'caption': ('Figure 15. Temperature distribution (deg C) at breathing zone height Z = 1.50 m, '
                    'five time snapshots (t = 120, 240, 360, 480, 600 s). Colour fill: inferno scale '
                    '(ambient 20 deg C to peak). Contour lines at 50, 100, 200, 400 deg C. '
                    'The thermal plume from CAB-4 propagates outward; buoyancy drives hot gases '
                    'upward so breathing-zone temperatures remain moderate beyond the immediate '
                    'cabinet bay. EQIX SG4-4A BESS Room 1, 9 ACH Stage 0 ventilation.'),
    },
    'fig15b': {
        'file':    'Temp_Ceiling_PBZ270.png',
        'title':   'Figure 15b: FDS CFD Simulation -- Temperature at Ceiling (Z = 2.70 m, Plan View)',
        'caption': ('Figure 15b. Temperature distribution (deg C) at ceiling height Z = 2.70 m, '
                    'five time snapshots (t = 120, 240, 360, 480, 600 s). Colour fill: inferno scale. '
                    'Contour lines at 50, 100, 200, 400 deg C. '
                    'Hot ceiling jet from CAB-4 fire (X=2.90-3.70) spreads along ceiling; '
                    'peak temperatures concentrated above the fire source and exhaust grilles. '
                    'EQIX SG4-4A BESS Room 1, 9 ACH Stage 0 ventilation.'),
    },
    'fig16': {
        'file':    'HF_Ceiling_PBZ270.png',
        'title':   'Figure 16: FDS CFD Simulation -- HF Concentration at Ceiling (Z = 2.70 m)',
        'caption': ('Figure 16. HF mass fraction at ceiling level Z = 2.70 m, five time snapshots '
                    '(t = 120, 240, 360, 480, 600 s). Red contour = IDLH-equivalent mass fraction '
                    '(30 ppm). HF accumulates at ceiling early; peak concentration near exhaust '
                    'grilles confirms ceiling-jet transport. '
                    'EQIX SG4-4A BESS Room 1, 9 ACH Stage 0 ventilation, NMC Galaxy LBF 242.76 kWh.'),
    },
    'fig17': {
        'file':    'H2_Ceiling_PBZ270.png',
        'title':   'Figure 17: FDS CFD Simulation -- H2 Concentration at Ceiling (Z = 2.70 m)',
        'caption': ('Figure 17. Hydrogen (H2) mass fraction at ceiling level Z = 2.70 m, five time '
                    'snapshots (t = 120, 240, 360, 480, 600 s). Red contour = 4% vol LFL '
                    '(0.00278 kg/kg). H2 is buoyant and accumulates at ceiling early; LFL contour '
                    'tracks outward from the source cabinet. Under 9 ACH ventilation, ceiling H2 '
                    'declines after t = 360 s. EQIX SG4-4A BESS Room 1.'),
    },
    'fig18': {
        'file':    'H2_BZ_Plan_PBZ150.png',
        'title':   'Figure 18: FDS CFD Simulation -- H2 Concentration at Breathing Zone (Z = 1.50 m)',
        'caption': ('Figure 18. H2 mass fraction at breathing zone height Z = 1.50 m, five time '
                    'snapshots. Orange contour = 25% LFL (explosion sensor alarm threshold, '
                    '0.000695 kg/kg); red contour = 100% LFL (0.00278 kg/kg). '
                    'H2 at breathing zone remains below the 25% LFL sensor threshold throughout '
                    '600 s, confirming that explosion risk at occupant height is low under '
                    '9 ACH ventilation. Ceiling-level H2 (Figure 17) is the primary explosion '
                    'concern. EQIX SG4-4A BESS Room 1.'),
    },
    'fig19': {
        'file':    'FireSection_X324.png',
        'title':   'Figure 19: FDS CFD Simulation -- Vertical Section Through CAB-4 (X = 3.24 m)',
        'caption': ('Figure 19. Vertical section at X = 3.24 m through CAB-4 thermal runaway source, '
                    'three rows: (top) HRRPUV -- heat release rate per unit volume (kW/m3); '
                    '(middle) temperature (deg C); (bottom) HF mass fraction (ppm). '
                    'Four time snapshots: t = 150, 300, 450, 600 s. '
                    'Dotted line at Z = 1.50 m (breathing zone). Fire plume peaks at t = 300 s; '
                    'HRRPUV declines as energy content exhausts. HF section confirms vertical '
                    'stratification -- ceiling jet visible. '
                    'EQIX SG4-4A BESS Room 1, 9 ACH Stage 0 ventilation.'),
    },
    'fig20': {
        'file':    'HF_Sensor_TimeHistory.png',
        'title':   'Figure 20: FDS CFD Simulation -- HF Sensor Time History',
        'caption': ('Figure 20. HF concentration time history at all FDS measurement devices over '
                    '600 s. Threshold lines: 0.5 ppm OSHA ceiling (blue dotted); '
                    '3 ppm ACGIH STEL (orange dashed); 30 ppm NIOSH IDLH (red solid). '
                    'TR initiates at t = 60 s. Peak sensor readings near t = 360-400 s; '
                    '9 ACH ventilation progressively reduces concentration after peak. '
                    'All devices exceed IDLH from approximately t = 120 s, consistent with '
                    'the well-mixed box model (Section 4.1). '
                    'EQIX SG4-4A BESS Room 1, NMC Galaxy LBF 242.76 kWh.'),
    },
}


# =============================================================================
def main():
    doc = Document(MANUSCRIPT)
    paras = doc.paragraphs
    print(f'Loaded {len(paras)} paragraphs from v4')

    # --- Collect element references BEFORE any modifications ---
    # (so index shifts from insertions do not matter)
    ref = {}
    search_map = {
        'p159': 'Critically: FM Global is quantitatively correct',
        'p178': '4.7 CFD Proxy Model',
        'p187': None,   # blank image para -- identified by preceding content
        'p309': None,   # supplementary Fig 12 image
    }

    # Find by content
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if 'Critically: FM Global is quantitatively correct' in txt:
            ref['p159'] = p
            print(f'  p159 (Section 4.5.4 end) at index {i}: {txt[:60]}')
        elif '4.7 CFD Proxy Model' in txt:
            ref['p178'] = p
            print(f'  p178 (Sec 4.7 heading) at index {i}: {txt[:60]}')
        elif 'Figure 14: CFD Proxy Model' in txt or 'Figure 14: FDS CFD' in txt:
            ref['p188'] = p
            print(f'  p188 (Fig14 title) at index {i}: {txt[:60]}')
        elif txt.startswith('Figure 14.') and 'CFD proxy' in txt.lower():
            ref['p189'] = p
            print(f'  p189 (Fig14 caption) at index {i}: {txt[:60]}')
        elif 'Figure 12: CFD-Analytical Dual Suppression' in txt:
            ref['p310'] = p
            print(f'  p310 (Supp Fig12 title) at index {i}: {txt[:60]}')
        elif txt.startswith('Figure 12: Three-phase CFD'):
            ref['p311'] = p
            print(f'  p311 (Supp Fig12 caption) at index {i}: {txt[:60]}')

    # Identify blank image paragraphs by finding them between known content
    # Para 187 (Fig 14 image) is right before the Fig 14 title (p188)
    # Para 309 (Supp Fig 12 image) is right before the Supp Fig 12 title (p310)
    if 'p188' in ref:
        # Derive p187 (image, one before) and p189 (caption body, one after) from p188
        all_p = doc.paragraphs
        idx188 = next(i for i, p in enumerate(all_p) if p._p is ref['p188']._p)
        ref['p187'] = all_p[idx188 - 1]
        ref['p189'] = all_p[idx188 + 1]
        pics = ref['p187']._p.findall('.//' + qn('a:blip'))
        print(f'  p187 (Fig14 image) at index {idx188-1}: '
              f'{"IMAGE OK" if pics else "NO IMAGE -- check"}')
        print(f'  p189 (Fig14 caption) at index {idx188+1}: '
              f'{ref["p189"].text.strip()[:60]}')

    if 'p310' in ref:
        all_p = doc.paragraphs
        idx310 = next(i for i, p in enumerate(all_p) if p._p is ref['p310']._p)
        ref['p309'] = all_p[idx310 - 1]
        pics = ref['p309']._p.findall('.//' + qn('a:blip'))
        print(f'  p309 (Supp Fig12 image) at index {idx310-1}: '
              f'{"IMAGE OK" if pics else "NO IMAGE -- check"}')

    missing = [k for k in ['p159', 'p178', 'p188', 'p189', 'p187',
                            'p309', 'p310', 'p311'] if k not in ref]
    if missing:
        print(f'ERROR: could not locate paragraphs: {missing}')
        return

    # =========================================================================
    # 1. Update Section 4.7 heading (now FDS results, not proxy model)
    # =========================================================================
    set_plain(ref['p178'],
              '4.7 FDS CFD Simulation: Spatial Dispersion Results')
    print('\n[1] Updated Section 4.7 heading')

    # =========================================================================
    # 2. Replace Figure 14 image and captions in main body
    # =========================================================================
    img14 = os.path.join(FIGS_DIR, FIGS['fig14']['file'])
    replace_with_image(ref['p187'], img14)
    set_bold(ref['p188'], FIGS['fig14']['title'])
    set_italic(ref['p189'], FIGS['fig14']['caption'])
    print('[2] Replaced Figure 14 image and updated captions')

    # =========================================================================
    # 3. Insert Figure 12 after Section 4.5.4 closing text (para 159)
    #    Order of insertions: each call inserts AFTER the previous result,
    #    so call insert_figure_block_after once.
    # =========================================================================
    img12 = os.path.join(FIGS_DIR, FIGS['fig12']['file'])
    insert_figure_block_after(ref['p159'], img12,
                               FIGS['fig12']['title'], FIGS['fig12']['caption'])
    print('[3] Inserted Figure 12 into Section 4.5.4')

    # =========================================================================
    # 4. Insert Figures 15-20 into Section 4.7 after Figure 14 caption
    #    IMPORTANT: insertions after p189 -- use chaining so each figure
    #    is inserted after the previous one in correct order.
    # =========================================================================
    fig_order = ['fig15', 'fig15b', 'fig16', 'fig17', 'fig18', 'fig19', 'fig20']
    last_para = ref['p189']
    for key in fig_order:
        fig_info = FIGS[key]
        img_path = os.path.join(FIGS_DIR, fig_info['file'])
        if not os.path.exists(img_path):
            print(f'  WARNING: {img_path} not found -- skipping {key}')
            continue
        last_para = insert_figure_block_after(
            last_para, img_path, fig_info['title'], fig_info['caption'])
        print(f'[4] Inserted {key}: {fig_info["file"]}')

    # =========================================================================
    # 5. Supplementary Figure 12: remove embedded image, update captions
    #    so the supplementary cross-references the main body figure
    # =========================================================================
    clear_runs(ref['p309'])   # remove image; paragraph becomes blank
    set_bold(ref['p310'],
             'Figure 12: See Section 4.5.4 in main paper')
    set_italic(ref['p311'],
               'Figure 12 (FDS CFD centreline section, Y = 2.50 m) is presented '
               'in the main body of the paper at Section 4.5.4. See Figure 12 above.')
    print('[5] Updated supplementary Figure 12 cross-reference')

    # =========================================================================
    doc.save(OUTPUT)
    print(f'\nSaved: {OUTPUT}')
    final_count = len(doc.paragraphs)
    print(f'Final paragraph count: {final_count}')


if __name__ == '__main__':
    main()
