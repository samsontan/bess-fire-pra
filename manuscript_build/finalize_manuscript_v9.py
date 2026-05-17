"""
finalize_manuscript_v9.py

Applies all changes from Sam's tracked-changes review to produce v9.docx:

1. Fix inline body text (para 142, 156, 172, 185) - figure number updates
2. White-out 'Figure 12:' and 'Figure 13 --' title text from embedded FDS images
3. Replace event tree image (image1.png) with rebuilt v4 PNG
4. Save as BESS_Fire_Safety_Paper_Q1_revised_v9.docx

Source: v9_base.docx (tracked changes accepted via Word COM)
"""

import sys, zipfile, shutil, os, io, re
sys.stdout.reconfigure(encoding='utf-8')

from PIL import Image, ImageDraw
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

BASE = r'G:\My Drive\SAI\Projects\EQIX_SG4-4A_NFPA855_HMA_Waiver_Report\__Mickey7_BESS_Fire_Safety_Paper'
SRC  = os.path.join(BASE, 'BESS_Fire_Safety_Paper_Q1_revised_v9_base.docx')
OUT  = os.path.join(BASE, 'BESS_Fire_Safety_Paper_Q1_revised_v9.docx')
ET_V4 = os.path.join(BASE, 'Figure1_EventTree_v4.png')

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Fix inline body text figure references
# ─────────────────────────────────────────────────────────────────────────────
print('STEP 1: Fixing inline body text ...')

shutil.copy2(SRC, OUT)
doc = Document(OUT)

def replace_in_para(para, old, new):
    """Replace old->new across all runs in a paragraph, handling cross-run splits."""
    full = para.text
    if old not in full:
        return False
    # Try single-run replacement first
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Cross-run fallback: rebuild first run, clear rest
    if old in full:
        new_full = full.replace(old, new)
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ''
        return True
    return False

INLINE_FIXES = [
    # Para 142: remove "Figure 2 (panels A-D) presents the event tree comparison
    # and ERL results. The key branches are:" -> replace reference to deleted figure
    (
        'Figure 2 (panels A–D) presents the event tree comparison and ERL results. The key branches are:',
        'The event tree comparison (water-only vs gas+water) and ERL results are as follows. The key branches are:'
    ),
    # Para 156: old Figure 3 (FDS HF+Temp section) is now Figure 2
    ('Figure 3 presents', 'Figure 2 presents'),
    # Para 172: old Figure 4 (F-N curves) is now Figure 3
    ('Figure 4', 'Figure 3'),
    # Para 185: old Figure 5 is now Figure 4
    ('Figure 5', 'Figure 4'),
]

IS_CAPTION = re.compile(r'^Figure \d+[:\.]', re.IGNORECASE)

changes = 0
for para in doc.paragraphs:
    # Never touch caption lines — they were already correct from tracked changes acceptance
    if IS_CAPTION.match(para.text.strip()):
        continue
    for old, new in INLINE_FIXES:
        if replace_in_para(para, old, new):
            print(f'  Fixed: {old[:60]!r} -> {new[:50]!r}')
            changes += 1
            break  # one fix per paragraph max in this pass

doc.save(OUT)
print(f'  {changes} inline text changes applied.')

# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: White-out 'Figure 12:' from FDS HF+Temp image (image2.png)
#         White-out 'Figure 13 --' from F-N curves image (image3.png)
#         Replace event tree image (image1.png) with v4 PNG
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 2: Processing embedded images ...')

def whiteout_top_strip(img_bytes, strip_height_fraction=0.032):
    """White-out the top strip of an image (title line with figure number)."""
    img = Image.open(io.BytesIO(img_bytes)).convert('RGBA')
    w, h = img.size
    strip_px = int(h * strip_height_fraction)
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, w, strip_px], fill=(255, 255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()

# Read the v4 event tree PNG
with open(ET_V4, 'rb') as f:
    et_v4_bytes = f.read()

# Process the docx zip: modify image files in-place
import tempfile
tmp = OUT + '.tmp.zip'

with zipfile.ZipFile(OUT, 'r') as zin:
    with zipfile.ZipFile(tmp, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == 'word/media/image1.png':
                # Replace event tree with v4
                data = et_v4_bytes
                print(f'  Replaced: image1.png (event tree v4, {len(data):,} bytes)')

            elif item.filename == 'word/media/image2.png':
                # White-out "Figure 12:" title from FDS HF+Temp figure
                data = whiteout_top_strip(data, strip_height_fraction=0.030)
                print(f'  Whited-out title strip: image2.png (FDS HF+Temp, {len(data):,} bytes)')

            elif item.filename == 'word/media/image3.png':
                # White-out "Figure 13 --" title from F-N curves figure
                data = whiteout_top_strip(data, strip_height_fraction=0.038)
                print(f'  Whited-out title strip: image3.png (F-N curves, {len(data):,} bytes)')

            zout.writestr(item, data)

os.replace(tmp, OUT)
print(f'  Images updated in: {OUT}')

# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Verify
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 3: Verification ...')
doc2 = Document(OUT)
ref_start2 = next((i for i, p in enumerate(doc2.paragraphs) if p.text.strip() == 'References'), 999)

# Check no stale Figure 2 inline refs remain in body (excluding captions)
stale = []
for i, p in enumerate(doc2.paragraphs[:ref_start2]):
    t = p.text
    if 'Figure 2' in t and not re.match(r'^Figure 2', t.strip()):
        stale.append(f'[{i}] {t[:100]}')
if stale:
    print(f'  WARNING: stale "Figure 2" inline refs: {stale}')
else:
    print('  OK: no stale "Figure 2" inline body refs')

# Check all captions
print('  Figure captions:')
for p in doc2.paragraphs[:ref_start2]:
    if re.match(r'^Figure \d+:', p.text.strip()):
        print(f'    {p.text.strip()[:90]}')

print(f'\nSaved: {OUT}')
print(f'Size:  {os.path.getsize(OUT):,} bytes')
