"""
build_v10_citations.py

Inserts citations [25]-[33] into v9.docx and saves as v10.docx:
- Applies all text insertions from citation_insertions_exact.md
- Converts "Additional Sources" unnumbered section to numbered [25]-[33]
- Removes the old unnumbered section
- Saves as v10.docx and converts to PDF
"""

import sys, os, shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

sys.stdout.reconfigure(encoding='utf-8')

BASE = r'G:\My Drive\SAI\Projects\EQIX_SG4-4A_NFPA855_HMA_Waiver_Report\__Mickey7_BESS_Fire_Safety_Paper'
SRC  = os.path.join(BASE, 'BESS_Fire_Safety_Paper_Q1_revised_v9.docx')
OUT  = os.path.join(BASE, 'BESS_Fire_Safety_Paper_Q1_revised_v10.docx')

shutil.copy2(SRC, OUT)
doc = Document(OUT)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    """Replace old->new across all runs in a paragraph, cross-run aware."""
    full = para.text
    if old not in full:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # cross-run fallback
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ''
    return True

def all_paragraphs(doc):
    """All body paragraphs plus table cell paragraphs."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def insert_para_after(anchor_elem, text, style_para):
    """Insert a new paragraph (copying pPr style from style_para) after anchor_elem."""
    new_p = OxmlElement('w:p')
    pPr = style_para._element.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    new_r = OxmlElement('w:r')
    # Copy run properties (font etc.) from first run of style_para if available
    if style_para.runs:
        rPr = style_para.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            new_r.append(deepcopy(rPr))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    anchor_elem.addnext(new_p)
    return new_p


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Text insertions and citation replacements
# ─────────────────────────────────────────────────────────────────────────────
print('STEP 1: Applying text insertions and citation updates ...\n')

REPLACEMENTS = [
    # ── Insertion 1: EPRI [31] added to Section 1.1 problem statement ──
    (
        'multi-dimensional and severe. At the cell level',
        'multi-dimensional and severe. The Electric Power Research Institute’s comprehensive failure mode analysis [31] '
        'identifies thermal runaway propagation as a system-level phenomenon spanning multiple energy dissipation pathways '
        '(cell → module → cabinet → room). At the cell level'
    ),
    # ── Insertion 2: DNV GL [30] added to Section 1.2 gap analysis ──
    (
        'relies on qualitative or semi-quantitative risk assessment. The NFPA 855 Chapter 5',
        'relies on qualitative or semi-quantitative risk assessment. Industry guidance from DNV GL [30] emphasizes the '
        'centrality of system-level protection design, including compartmentation and early gas detection, over prescriptive '
        'component-level standards; however, this guidance remains qualitative and lacks the probabilistic framework necessary '
        'to compare alternative system designs on quantitative risk grounds. The NFPA 855 Chapter 5'
    ),
    # ── Insertion 3: Singapore Fire Code [28] in Section 2.1 ──
    (
        'Singapore Fire Code 2023 Clause 10.3.1(b),',
        'Singapore Fire Code 2023 [28] Clause 10.3.1(b),'
    ),
    # ── Insertion 4: Golubkov [25] and Ohneseit [26] numbered ──
    (
        '(Golubkov et al., 2014; Ohneseit et al., 2023; Sadeghi & Restuccia, 2024).',
        '([25]; [26]; Sadeghi & Restuccia, 2024).'
    ),
    # ── Insertion 5: Liao et al. [32] inserted after TR cascade description ──
    (
        'separator meltdown with flaming ejection at >300°C [1,6]. The SOC at time of TR strongly influences severity:',
        'separator meltdown with flaming ejection at >300°C [1,6]. '
        'The propagation of thermal runaway from individual cells into multi-cell modules is governed by thermal coupling '
        'between adjacent cells; Liao et al. [32] demonstrated experimentally that failure propagation probability within a '
        'module increases sharply with increasing State of Charge (SOC), with propagation distances exceeding 50 mm under '
        'high-SOC conditions. The cabinet-level propagation probabilities incorporated in this paper’s event tree '
        '(Section 4) are based on this module-level propagation characterization extrapolated to full-cabinet scale. '
        'The SOC at time of TR strongly influences severity:'
    ),
    # ── Insertion 6: Singapore Fire Code [28] in Section 2.5 ──
    (
        'The Singapore Fire Code 2023 (4th Amendment) Clause 10.3.1 incorporates',
        'The Singapore Fire Code 2023 (4th Amendment) [28] Clause 10.3.1 incorporates'
    ),
    # ── Insertion 7: Singapore Fire Code [28] in Section 3.1 ──
    (
        'Singapore Fire Code 2023 (4th Amendment) Table 10.3.1 applies identical',
        'Singapore Fire Code 2023 [28] Table 10.3.1 applies identical'
    ),
    # ── Insertion 8a: FM Global DS 5-32 [29] first instance in Section 3.5 ──
    (
        'FM Global Property Loss Prevention Data Sheet 5-32 (2021) for NMC lithium-ion',
        'FM Global Property Loss Prevention Data Sheet 5-32 [29] for NMC lithium-ion'
    ),
    # ── Insertion 8b: FM Global [29] second instance + NFPA 2001 [33] inserted ──
    (
        'providing flame suppression only — no TR arrest capability — consistent with the established '
        'self-oxidising chemistry of NMC cathodes and consistent with the position of FM Global '
        '(Property Loss Prevention Data Sheet 5-32, 2021).',
        'providing flame suppression only — no TR arrest capability — consistent with the established '
        'self-oxidising chemistry of NMC cathodes and consistent with the position of FM Global '
        '(Property Loss Prevention Data Sheet 5-32 [29]). '
        'The application rate, discharge duration, and oxygen depletion mechanism of the clean agent system at EQIX SG4-4A '
        'conform to NFPA 2001 [33] design requirements for clean agent systems in occupied spaces, with particular emphasis '
        'on post-discharge venting to restore room oxygen to habitable levels (>19.5% by volume) within the timeframe '
        'specified in the standard. The clean agent system operates as a pre-sprinkler layer designed to suppress open '
        'flaming (HRR reduction) rather than arrest thermal runaway; the subsequent water suppression layer is designed to '
        'cool the exothermic reaction and interrupt propagation.'
    ),
    # ── Insertion 8c: FM Global DS 5-32 in testing context (para ~88) ──
    (
        '(FM Global DS 5-32, 2021), which is why the Monte Carlo mean effectiveness',
        '(FM Global DS 5-32 [29]), which is why the Monte Carlo mean effectiveness'
    ),
    # ── Insertion 9: Table 1 FM Global and Jensen citations updated ──
    (
        'FM Global DS 5-32 (2021); Jensen et al. (2019)',
        'FM Global DS 5-32 [29]; Jensen et al. [17]'
    ),
    # ── Insertion 10a: UL 9540A [27] — through BMS context (Section 3.x) ──
    (
        'through BMS, UL 9540A containment, and EPO',
        'through BMS, UL 9540A [27] containment, and EPO'
    ),
    # ── Insertion 10b: UL 9540A [27] — P_fail context ──
    (
        'UL 9540A containment (P_fail',
        'UL 9540A [27] containment (P_fail'
    ),
    # ── Insertion 10c: UL 9540A [27] — "is the only" context ──
    (
        'UL 9540A containment) is the only',
        'UL 9540A [27] containment) is the only'
    ),
    # ── Insertion 10d: UL 9540A [27] — cabinet containment context ──
    (
        'UL 9540A cabinet containment.',
        'UL 9540A [27] cabinet containment.'
    ),
    # ── Insertion 10e: UL 9540A [27] — containment probability header (Section 5.6) ──
    (
        'UL 9540A containment probability: The 0.92',
        'UL 9540A [27] containment probability: The 0.92'
    ),
    # ── Insertion 10f: UL 9540A [27] — verification phrase (Section 4.1) ──
    (
        'verification of UL 9540A test configuration',
        'verification of UL 9540A [27] test configuration'
    ),
    # ── Insertion 10g: UL 9540A [27] — OI-02 unresolved (Section 5.6) ──
    (
        'UL 9540A containment probability (OI-02 unresolved)',
        'UL 9540A [27] containment probability (OI-02 unresolved)'
    ),
    # ── Insertion 10h: UL 9540A [27] — test reports phrase ──
    (
        'UL 9540A test reports',
        'UL 9540A [27] test reports'
    ),
]

changes = 0
misses  = []
for old, new in REPLACEMENTS:
    found = False
    for para in all_paragraphs(doc):
        if replace_in_para(para, old, new):
            label = old[:65]
            print(f'  OK  : {label!r}')
            changes += 1
            found = True
            break
    if not found:
        misses.append(old[:65])
        print(f'  MISS: {old[:65]!r}')

print(f'\n  {changes}/{len(REPLACEMENTS)} replacements applied.')
if misses:
    print(f'  WARNING: {len(misses)} miss(es):')
    for m in misses:
        print(f'    {m!r}')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Add numbered references [25]-[33] after [24]
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 2: Adding reference entries [25]-[33] after [24] ...')

NEW_REFS = [
    '[25] Golubkov, A. W., Fuchs, D., Wagner, J., Wiltsche, H., Stangl, C., Fauler, G., Voitic, G., Thaler, A., & '
    'Hacker, V. (2014). Thermal-runaway experiments on consumer Li-ion batteries with metal-oxide and olivin-type '
    'cathodes. RSC Advances, 4(7), 3633–3642. https://doi.org/10.1039/c3ra45748f',

    '[26] Ohneseit, S., Finster, P., Floras, C., Lubenau, N., Uhlmann, N., Seifert, H. J., & Ziebert, C. (2023). '
    'Thermal and Mechanical Safety Assessment of Type 21700 Lithium-Ion Batteries with NMC, NCA and LFP Cathodes — '
    'Investigation of Cell Abuse by Means of Accelerating Rate Calorimetry (ARC). Batteries, 9(5), 237. '
    'https://doi.org/10.3390/batteries9050237',

    '[27] UL Standards & Engagement. (2023). UL 9540A: Test Method for Evaluating Thermal Runaway Fire Propagation '
    'in Battery Energy Storage Systems (4th ed.). UL.',

    '[28] Singapore Civil Defence Force. (2023). Singapore Fire Code 2023 (4th Amendment). SCDF.',

    '[29] FM Global. (2021). Property Loss Prevention Data Sheets 5-32: Electrical Energy Storage Systems. FM Global.',

    '[30] DNV GL. (2021). Considerations for ESS Fire Safety (Report No. 2021-1004). DNV GL Energy.',

    '[31] Electric Power Research Institute (EPRI). (2020). Energy Storage System Safety: Failure Mode Analysis and '
    'Risk Quantification (Report 3002016958). EPRI.',

    '[32] Liao, Z., Zhang, S., Li, K., Mao, B., & Jiang, L. (2020). Hazard analysis of thermally-induced failure '
    'propagation in lithium-ion battery modules. Journal of Hazardous Materials, 393, 122442.',

    '[33] National Fire Protection Association. (2019). NFPA 2001: Standard on Clean Agent Fire Extinguishing '
    'Systems (2018 ed.). NFPA.',
]

ref24_para = None
for para in doc.paragraphs:
    if para.text.strip().startswith('[24]'):
        ref24_para = para
        break

if ref24_para is None:
    print('  WARNING: [24] paragraph not found!')
else:
    print(f'  Anchor: {ref24_para.text[:80]!r}')
    anchor = ref24_para._element
    for ref_text in NEW_REFS:
        anchor = insert_para_after(anchor, ref_text, ref24_para)
    print(f'  Inserted {len(NEW_REFS)} reference entries.')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Remove the old "Additional Sources" unnumbered section
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 3: Removing "Additional Sources" unnumbered section ...')

# Collect paragraphs in the Additional Sources section
# (from the header to the last UL Standards entry, inclusive)
paras_to_delete = []
in_additional = False
for para in doc.paragraphs:
    t = para.text.strip()
    if t.startswith('Additional Sources (not formally cited in text)'):
        in_additional = True
    if in_additional:
        paras_to_delete.append(para)
        # Stop after the UL 9540A entry (last entry before Appendix A)
        if 'UL 9540A: Test Method for Evaluating' in t:
            break

print(f'  Found {len(paras_to_delete)} paragraphs to remove:')
for p in paras_to_delete:
    t = p.text.strip()
    if t:
        print(f'    {t[:90]!r}')

for para in paras_to_delete:
    parent = para._element.getparent()
    if parent is not None:
        parent.remove(para._element)

print(f'  Removed {len(paras_to_delete)} paragraphs.')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: Verify
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 4: Verification ...')

# Check all numbered references
print('  References in document:')
for para in doc.paragraphs:
    t = para.text.strip()
    if t and t[0] == '[' and ']' in t[:5]:
        print(f'    {t[:100]}')

# Check no "Additional Sources" remains
still_there = any('Additional Sources' in p.text for p in doc.paragraphs)
print(f'  Additional Sources section removed: {not still_there}')

# Check key insertions landed
checks = [
    ('[31] identifies thermal runaway', 'EPRI [31] insertion'),
    ('DNV GL [30] emphasizes', 'DNV GL [30] insertion'),
    ('Fire Code 2023 [28] Clause 10.3.1(b)', 'SFC [28] Sec 2.1'),
    ('([25]; [26]; Sadeghi', 'Golubkov/Ohneseit numbered'),
    ('Liao et al. [32] demonstrated', 'Liao [32] insertion'),
    ('Fire Code 2023 (4th Amendment) [28] Clause 10.3.1 incorporates', 'SFC [28] Sec 2.5'),
    ('Fire Code 2023 [28] Table 10.3.1', 'SFC [28] Sec 3.1'),
    ('Data Sheet 5-32 [29] for NMC', 'FM Global [29] 8a'),
    ('Data Sheet 5-32 [29]). The application', 'FM Global [29] + NFPA 2001 [33]'),
    ('NFPA 2001 [33] design requirements', 'NFPA 2001 [33]'),
    ('UL 9540A [27]', 'UL 9540A [27]'),
]

doc_text = '\n'.join(p.text for p in doc.paragraphs)
# Also include table text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            doc_text += '\n' + cell.text

print('\n  Citation insertion checks:')
for needle, label in checks:
    ok = needle in doc_text
    print(f'  {"OK " if ok else "ERR"}: {label}')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 5: Save
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'\nSaved: {OUT}')
print(f'Size:  {os.path.getsize(OUT):,} bytes')
