"""
prepare_bess_draft_for_reviewers.py

1. Update v10.docx author block to 4-author format (XAI paper style)
2. Remove ORCID and Correspondence lines
3. Save as clean-named draft + PDF
4. Copy to "Manuscript Draft to Khalid and Paul - 2026-05-17" folder
5. Create email draft document
"""

import sys, os, shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from copy import deepcopy

sys.stdout.reconfigure(encoding='utf-8')

BASE    = r'G:\My Drive\SAI\Projects\EQIX_SG4-4A_NFPA855_HMA_Waiver_Report\__Mickey7_BESS_Fire_Safety_Paper'
SRC_V10 = os.path.join(BASE, 'BESS_Fire_Safety_Paper_Q1_revised_v10.docx')

DATE    = '2026-05-17'
FOLDER  = os.path.join(BASE, f'Manuscript Draft to Khalid and Paul - {DATE}')
DRAFT_DOCX = os.path.join(FOLDER, f'BESS_Fire_Safety_Paper_Draft_{DATE}.docx')
DRAFT_PDF  = os.path.join(FOLDER, f'BESS_Fire_Safety_Paper_Draft_{DATE}.pdf')
EMAIL_DOCX = os.path.join(FOLDER, f'Email Draft - Khalid and Paul.docx')

os.makedirs(FOLDER, exist_ok=True)
shutil.copy2(SRC_V10, DRAFT_DOCX)

doc = Document(DRAFT_DOCX)


# ─────────────────────────────────────────────────────────────────────────────
# Helper: clear runs from a paragraph and rebuild
# ─────────────────────────────────────────────────────────────────────────────
def clear_runs(para):
    p_elem = para._element
    for child in list(p_elem):
        if child.tag in (qn('w:r'), qn('w:hyperlink'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')):
            p_elem.remove(child)

def add_run(para, text, superscript=False, italic=False, bold=False):
    run = para.add_run(text)
    run.font.superscript = superscript
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    return run

def insert_para_after(anchor_elem, style_para):
    """Insert empty paragraph (copying pPr from style_para) after anchor_elem. Returns element."""
    new_p = OxmlElement('w:p')
    pPr = style_para._element.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    anchor_elem.addnext(new_p)
    return new_p


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Locate author block paragraphs by scanning text
# ─────────────────────────────────────────────────────────────────────────────
print('STEP 1: Locating author block ...')

paras = doc.paragraphs
author_idx   = None
affil1_idx   = None
affil2_idx   = None

for i, para in enumerate(paras[:15]):
    t = para.text.strip()
    if 'Samson' in t and 'Tan' in t and author_idx is None:
        author_idx = i
        print(f'  Author line:    [{i}] {t[:60]!r}')
    elif author_idx is not None and affil1_idx is None and i == author_idx + 1:
        affil1_idx = i
        print(f'  Affiliation 1:  [{i}] {t[:60]!r}')
    elif affil1_idx is not None and affil2_idx is None and i == affil1_idx + 1:
        affil2_idx = i
        print(f'  Affiliation 2:  [{i}] {t[:60]!r}')

author_para = paras[author_idx] if author_idx is not None else None
affil1_para = paras[affil1_idx] if affil1_idx is not None else None
affil2_para = paras[affil2_idx] if affil2_idx is not None else None


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Rebuild author line with 4 authors and superscript affiliations
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 2: Rebuilding author line ...')

if author_para:
    clear_runs(author_para)
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(author_para, 'Samson Tan')
    add_run(author_para, '1,3', superscript=True)
    add_run(author_para, ', Teoh Teik Toe')
    add_run(author_para, '2,3', superscript=True)
    add_run(author_para, ', Paul Joseph')
    add_run(author_para, '1', superscript=True)
    add_run(author_para, ', Khalid Moinuddin')
    add_run(author_para, '1', superscript=True)
    print(f'  Author line set: {author_para.text!r}')
else:
    print('  ERROR: author paragraph not found!')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Rebuild affiliation lines (italic, centered)
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 3: Rebuilding affiliation lines ...')

AFFIL_1 = '1 Victoria University, Melbourne, Institute for Sustainable Industries and Liveable Cities, Australia'
AFFIL_2 = '2 Nanyang Technological University, Artificial Intelligence, Nanyang Business School, Singapore'
AFFIL_3 = '3 Staarch Pte Ltd Singapore'

if affil1_para:
    clear_runs(affil1_para)
    affil1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(affil1_para, AFFIL_1, italic=True)
    print(f'  Affil 1: {affil1_para.text[:70]!r}')

if affil2_para:
    clear_runs(affil2_para)
    affil2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(affil2_para, AFFIL_2, italic=True)
    print(f'  Affil 2: {affil2_para.text[:70]!r}')

# Insert affiliation 3 after affil2
if affil2_para:
    new_p_elem = insert_para_after(affil2_para._element, affil2_para)
    # Build a run inside it
    new_r = OxmlElement('w:r')
    # Copy rPr (run props) from first run of affil2 to get italic
    if affil2_para.runs:
        rPr = affil2_para.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            new_r.append(deepcopy(rPr))
    new_t = OxmlElement('w:t')
    new_t.text = AFFIL_3
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p_elem.append(new_r)
    print(f'  Affil 3: inserted {AFFIL_3!r}')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: Delete ORCID and Correspondence paragraphs
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 4: Removing ORCID and Correspondence paragraphs ...')

# Re-scan to catch any by text (in case indices shifted)
to_delete = []
for para in doc.paragraphs[:20]:
    t = para.text.strip()
    if t.startswith('ORCID') or t.startswith('Correspondence:'):
        to_delete.append(para)

for para in to_delete:
    print(f'  Deleting: {para.text[:60]!r}')
    parent = para._element.getparent()
    if parent is not None:
        parent.remove(para._element)

if not to_delete:
    print('  No ORCID/Correspondence paragraphs found (may already be clean).')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 5: Save draft docx
# ─────────────────────────────────────────────────────────────────────────────
doc.save(DRAFT_DOCX)
print(f'\nSTEP 5: Draft saved: {DRAFT_DOCX}')
print(f'  Size: {os.path.getsize(DRAFT_DOCX):,} bytes')

# Verify author block
print('\n  First-page verification:')
doc2 = Document(DRAFT_DOCX)
for i, para in enumerate(doc2.paragraphs[:12]):
    t = para.text.strip()
    if t:
        print(f'  [{i}] {t[:100]!r}')


# ─────────────────────────────────────────────────────────────────────────────
# STEP 6: Convert to PDF via Word COM
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 6: Converting to PDF ...')
import win32com.client

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
try:
    wdoc = word.Documents.Open(os.path.abspath(DRAFT_DOCX))
    wdoc.SaveAs(os.path.abspath(DRAFT_PDF), FileFormat=17)
    wdoc.Close(False)
    print(f'  PDF: {DRAFT_PDF}')
    print(f'  Size: {os.path.getsize(DRAFT_PDF):,} bytes')
finally:
    word.Quit()


# ─────────────────────────────────────────────────────────────────────────────
# STEP 7: Create email draft document
# ─────────────────────────────────────────────────────────────────────────────
print('\nSTEP 7: Creating email draft ...')

email_doc = Document()

# Remove default empty paragraph
for para in list(email_doc.paragraphs):
    p = para._element
    p.getparent().remove(p)

def ep(doc, text='', bold=False, italic=False):
    """Add a paragraph to the email doc."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return p

ep(email_doc, 'TO:', bold=True)
ep(email_doc, 'Professor Khalid Moinuddin (k.moinuddin@vu.edu.au)')
ep(email_doc, 'Associate Professor Paul Joseph (paul.joseph@vu.edu.au)')
ep(email_doc)
ep(email_doc, 'CC:', bold=True)
ep(email_doc, 'Teoh Teik Toe')
ep(email_doc)
ep(email_doc, 'SUBJECT:', bold=True)
ep(email_doc, 'Draft Manuscript for Review: Probabilistic Risk Assessment of BESS Fire Hazards')
ep(email_doc)
ep(email_doc, '---')
ep(email_doc)
ep(email_doc, 'Dear Khalid and Paul,')
ep(email_doc)
ep(email_doc,
   'Please find attached a draft manuscript titled "Probabilistic Risk Assessment of Grid-Scale '
   'Lithium-Ion Battery Energy Storage System Fire Hazards: A Monte Carlo Simulation Framework '
   'for Hydrogen Fluoride Toxicity and Suppression Effectiveness." We would very much appreciate '
   'your review and comments before we submit for journal consideration.')
ep(email_doc)
ep(email_doc, 'Background', bold=True)
ep(email_doc,
   'This paper arose directly from a Hazard Mitigation Analysis (HMA) report that STAARCH '
   'prepared for a 485.52 kWh NMC battery energy storage system installation at a data centre '
   'operated by a major multinational at 7 Tai Seng Avenue, Singapore. The installation is a '
   'Level 5 BESS under SCDF Fire Code classification and required SCDF prior approval via the '
   'NFPA 855 Exception (1) pathway.')
ep(email_doc)
ep(email_doc, 'Why this paper exists', bold=True)
ep(email_doc,
   'The client\'s installation incorporated two layers of fire suppression -- a clean agent '
   '(Fluoro-K / HFC-227ea) pre-action system followed by a water mist sprinkler system. This '
   'two-stage approach runs counter to FM Global Property Loss Prevention Data Sheet 5-32, which '
   'prescribes water as the only effective suppression medium for lithium-ion fires and does not '
   'recognise clean agent pre-discharge as contributing to life safety outcomes.')
ep(email_doc)
ep(email_doc,
   'Rather than accept the FM Global position at face value, we applied a full Monte Carlo '
   'probabilistic risk assessment (N = 10,000 iterations) to quantify the suppression '
   'effectiveness of each layer and the residual risk under both single-stage (water-only) and '
   'two-stage (gas + water) designs. The PRA demonstrates that the dual suppression system '
   'reduces annual Expected Risk to Life (ERL) by 80.3% -- from 1.22 x 10^-4 (water-only, '
   'ALARP-tolerable boundary) to 2.4 x 10^-5 (broadly acceptable). This provides a defensible '
   'quantitative justification for the two-stage design that the FM Global qualitative framework '
   'cannot supply.')
ep(email_doc)
ep(email_doc, 'The NMC vs LFP question', bold=True)
ep(email_doc,
   'A secondary driver for this paper is the chemistry justification. SCDF\'s informal position '
   'during the HMA review process favoured LFP (lithium iron phosphate) chemistry on the basis '
   'of its higher thermal runaway onset temperature (>270°C vs 130-150°C for NMC). However, the '
   'client\'s specified equipment -- 14 Schneider Electric Galaxy LBF cabinets -- uses NMC '
   'chemistry, and the project could not be redesigned around a chemistry substitution. '
   'The PRA framework in this paper provides the quantitative basis for demonstrating that, '
   'with the two-stage suppression design and two-compartment spatial arrangement, the residual '
   'risk of the NMC installation falls within the broadly acceptable ALARP band -- which is the '
   'substantive test, not chemistry selection per se.')
ep(email_doc)
ep(email_doc, 'What we are asking', bold=True)
ep(email_doc,
   'We are at review draft stage. We would welcome your comments on:')

for point in [
    'The event tree structure and probability assignments (Section 4)',
    'The suppression effectiveness model and Monte Carlo parameterisation (Section 3.5)',
    'The HF toxicity and ERL framework (Sections 2.3, 3.4)',
    'The framing of the FM Global / SCDF regulatory tension in the discussion (Section 5)',
    'Any references you feel are missing or where you have recent work we should cite',
]:
    p = email_doc.add_paragraph(style='List Bullet')
    p.add_run(point)

ep(email_doc)
ep(email_doc,
   'There is no deadline pressure -- please review at your convenience. We attach both the '
   'Word document and a PDF for ease of reading.')
ep(email_doc)
ep(email_doc, 'Thank you, Khalid and Paul. We look forward to your thoughts.')
ep(email_doc)
ep(email_doc, 'With regards,')
ep(email_doc, 'Samson Tan')
ep(email_doc, 'STAARCH Pte Ltd')
ep(email_doc, 'samson.tan@staarch.com')

email_doc.save(EMAIL_DOCX)
print(f'  Email draft: {EMAIL_DOCX}')
print(f'  Size: {os.path.getsize(EMAIL_DOCX):,} bytes')

print(f'\n{"="*60}')
print(f'All files in: {FOLDER}')
for f in os.listdir(FOLDER):
    sz = os.path.getsize(os.path.join(FOLDER, f))
    print(f'  {f:<50s}  {sz:>10,} bytes')
